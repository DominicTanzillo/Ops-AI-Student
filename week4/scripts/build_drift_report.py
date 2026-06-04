"""Builds the Drift Detection Report .docx skeleton with all figures + data
populated from the local JSON outputs. Leaves prose to the user.

Styling matches the user's Week 4 Report.docx:
  - No Word heading styles (plain paragraphs only)
  - No colors, no italics for commentary placeholders
  - Simple Table Grid only (no Accent styles)
  - ASCII characters only: no em dashes, no arrows, no checkmarks,
    no Greek letters, no multiplication signs
  - Section labels are bold paragraphs, not headings

Inputs (all in week4/data/, gitignored - read locally):
  level2_metrics.json   - prediction-level evidence (RMSE / deviance / MAPE)
  metrics_report.json   - 8-metric monitoring run output
  drift_patterns.json   - 4-pattern detector output

Output:
  week4/Week4_DriftDetectionReport.docx
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[2]
WEEK4 = REPO / "week4"
DATA = WEEK4 / "data"
OUT = WEEK4 / "Week4_DriftDetectionReport.docx"

BOROUGH_NAMES = {0: "Manhattan", 1: "Queens", 2: "Brooklyn", 3: "Bronx", 4: "Staten Island"}


def _load(name: str) -> dict:
    return json.loads((DATA / name).read_text())


def _section(doc, text: str) -> None:
    """Bold paragraph used as a section label (no Word heading style)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def _sub(doc, text: str) -> None:
    """Smaller bold paragraph used as a subsection label."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def _para(doc, text: str) -> None:
    doc.add_paragraph(text)


def _commentary(doc, hint: str) -> None:
    """Placeholder for student commentary. Plain text in brackets."""
    doc.add_paragraph(f"[YOUR COMMENTARY: {hint}]")


def _table(doc, headers: list[str], rows: list[list]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = _fmt(val)
    doc.add_paragraph()


def _fmt(v) -> str:
    if v is None:
        return "-"
    if isinstance(v, float):
        if abs(v) < 0.001 and v != 0:
            return f"{v:.2e}"
        return f"{v:.4f}"
    return str(v)


def build() -> None:
    level2 = _load("level2_metrics.json")
    metrics = _load("metrics_report.json")
    patterns = _load("drift_patterns.json")

    doc = Document()

    # ------------------------------------------------------------------
    # Title block (plain paragraphs)
    # ------------------------------------------------------------------
    _section(doc, "Week 4 Drift Detection Report")
    p = doc.add_paragraph()
    p.add_run("Student: ").bold = True
    p.add_run("Dominic Tanzillo (NetID: dpt7)")
    p = doc.add_paragraph()
    p.add_run("Course: ").bold = True
    p.add_run("AIPI 561 Operationalizing AI, Summer 2026")
    p = doc.add_paragraph()
    p.add_run("Window analyzed: ").bold = True
    p.add_run("Feb 2 to 28, 2026 (drift window) versus 2023-01-01 to 2026-01-15 (baseline)")
    p = doc.add_paragraph()
    p.add_run("Model: ").bold = True
    p.add_run("LightGBM Poisson, 814 trees, from Week 2 (week2/model/lgbm_demand_model.txt)")
    doc.add_paragraph()

    m1 = metrics["metrics"]["metric_1_accuracy"]

    # ------------------------------------------------------------------
    # Executive summary
    # ------------------------------------------------------------------
    _section(doc, "Executive Summary")
    _sub(doc, "Headline numbers (global, baseline vs Feb 2 to 28 drift window)")
    _table(doc,
        ["Metric", "Baseline", "Drift window", "Relative change", "Fires our threshold?"],
        [
            ["Poisson deviance",
             m1.get("baseline_poisson_deviance"),
             m1["poisson_deviance"],
             f"{m1['deviance_rel_delta']*100:+.1f}%",
             "CRITICAL (>= 20%)"],
            ["RMSE",
             m1["rmse"] / (1 + m1["rmse_rel_delta"]) if m1["rmse_rel_delta"] != 0 else m1["rmse"],
             m1["rmse"],
             f"{m1['rmse_rel_delta']*100:+.1f}%",
             "no (< 20%)"],
            ["MAPE (absolute change)",
             "-",
             m1["mape"],
             f"{m1['mape_abs_delta']:+.4f} pp",
             "not applicable"],
        ])
    _para(doc,
          f"Alert counts: {metrics['alerts']['critical_count']} CRITICAL, "
          f"{metrics['alerts']['warning_count']} WARNING. "
          f"Decision: {metrics['decision']}.")
    _commentary(doc,
                "1-paragraph TL;DR. State the verdict in plain English (e.g. "
                "the model is materially worse on Feb 2 to 28; 4 distinct drift "
                "patterns identified; retraining recommended). Cite the "
                "deviance versus RMSE contrast: it is your strongest single number.")

    # ------------------------------------------------------------------
    # Methodology
    # ------------------------------------------------------------------
    _section(doc, "Methodology")
    _para(doc, "Three-Level Investigation Framework.")
    _table(doc,
        ["Level", "What it inspects", "Script", "Output"],
        [
            ["Level 1 Feature",
             "Per-feature KS / PSI / Jensen-Shannon",
             "explore_drift_features.py",
             "console + manual review"],
            ["Level 2 Prediction",
             "Score both windows with the model; compare RMSE / deviance / MAPE",
             "level2_prediction_drift.py",
             "level2_metrics.json"],
            ["Level 3 Segment",
             "Per-zone / per-borough / per-hour / per (borough x is_weekend) cuts",
             "compute_metrics.py + detect_drift.py",
             "metrics_report.json + drift_patterns.json"],
        ])
    _para(doc,
          "The manifest file (week4_drift_manifest.json) was used as a "
          "self-check at the end only. No metric was designed by reading "
          "the manifest first.")
    _commentary(doc,
                "Optional: short paragraph on why the 3-level framework. "
                "Key point: feature-level catches data drift, prediction-level "
                "catches concept drift, segment-level localizes both.")

    # ------------------------------------------------------------------
    # Pattern 1
    # ------------------------------------------------------------------
    _section(doc, "Pattern 1: Temporal Peak Shift")
    _table(doc,
        ["Field", "Value"],
        [
            ["What drifted", "trip_count distribution per hour-of-day"],
            ["Type", "Data drift"],
            ["Detection layer", "metric_4 (KS test) per-hour cut"],
            ["Statistical test", "Two-sample Kolmogorov-Smirnov"],
        ])

    _sub(doc, "Worst hours by KS statistic")
    p1 = patterns["pattern_1_temporal_peak_shift"]
    rows = []
    for entry in p1["worst_5_hours_by_ks"]:
        h, stats = entry
        rows.append([h, stats["ks_statistic"], stats["baseline_mean"],
                     stats["new_mean"],
                     f"{(stats['new_mean']-stats['baseline_mean'])/max(stats['baseline_mean'],1e-9)*100:+.1f}%"])
    _table(doc,
        ["Hour", "KS statistic", "Baseline mean", "Drift mean", "Change in mean"],
        rows)

    _sub(doc, "Per-hour KS p-values (all 24 hours)")
    ks_per_hour = metrics["metrics"]["metric_4_ks_test"]["per_hour"]
    rows = []
    for h in sorted(ks_per_hour.keys(), key=int):
        s = ks_per_hour[h]
        sev = "***" if s["pvalue"] < 0.01 else ("**" if s["pvalue"] < 0.05 else "")
        rows.append([h, s["statistic"], s["pvalue"], sev])
    _table(doc,
        ["Hour", "KS statistic", "p-value", "Severity"],
        rows)

    _commentary(doc,
                "Explain what the table shows: which hours stand out, what "
                "the bimodal/temporal shape means physically (e.g. peak "
                "suppression in late-morning hours), and your root-cause "
                "hypothesis (employer schedule shift / event / pipeline).")

    # ------------------------------------------------------------------
    # Pattern 2
    # ------------------------------------------------------------------
    _section(doc, "Pattern 2: Manhattan Lag-Feature Deflation")
    _table(doc,
        ["Field", "Value"],
        [
            ["What drifted", "lag_1day, lag_1week, roll_mean_1day for Manhattan zones"],
            ["Type", "Data drift (input-feature)"],
            ["Detection layer", "metric_5 (PSI) per-borough cut"],
            ["Statistical test", "Population Stability Index (10 bins)"],
        ])

    _sub(doc, "PSI per borough on lag features")
    p2 = patterns["pattern_2_borough_feature_drift"]["per_borough"]
    rows = []
    for b in sorted(p2.keys(), key=int):
        s = p2[b]
        psi = s["psi"]
        rows.append([f"{b} ({s['borough_name']})",
                     psi.get("lag_1day", float("nan")),
                     psi.get("lag_1week", float("nan")),
                     psi.get("roll_mean_1day", float("nan")),
                     s.get("max_psi", float("nan"))])
    _table(doc,
        ["Borough", "lag_1day PSI", "lag_1week PSI", "roll_mean_1day PSI", "max PSI"],
        rows)
    _para(doc,
          "Reference thresholds: PSI < 0.10 means negligible. "
          "0.10 to 0.25 means small change (WARN). "
          "PSI > 0.25 means significant change (CRITICAL).")

    _commentary(doc,
                "Walk the reader through what the per-borough table reveals. "
                "Key observation: only one borough's lag features moved; the "
                "others stayed put. State your hypothesis about the upstream "
                "feature-engineering pipeline.")

    # ------------------------------------------------------------------
    # Pattern 3
    # ------------------------------------------------------------------
    _section(doc, "Pattern 3: Outer-Borough Baseline Scramble")
    _table(doc,
        ["Field", "Value"],
        [
            ["What drifted",
             "zone_slot_baseline to trip_count correlation for ~5 outer-borough zones"],
            ["Type",
             "Data drift (feature-to-target correlation destruction)"],
            ["Detection layer",
             "metric_2 per-borough deviance; per-zone actual-mean shift"],
        ])

    _sub(doc, "The Brooklyn paradox - per-borough Poisson deviance")
    pb = metrics["metrics"]["metric_2_accuracy_by_zone"]["per_borough"]
    rows = []
    for bid in sorted(pb.keys(), key=int):
        s = pb[bid]
        dev_delta = s.get("deviance_rel_delta")
        rows.append([
            f"{bid} ({BOROUGH_NAMES.get(int(bid), '?')})",
            s["n"],
            s["poisson_deviance"],
            f"{dev_delta*100:+.1f}%" if dev_delta is not None else "-",
            s.get("actual_mean"),
        ])
    _table(doc,
        ["Borough", "n rows", "Deviance (drift)", "Change vs baseline", "Actual mean"],
        rows)
    _para(doc,
          "Note: Brooklyn's deviance went DOWN, not up. This is the "
          "characteristic 'absolute error metric collapses when actuals "
          "collapse' artifact. Predictions are structurally wrong but the "
          "error looks small because the actuals also became small.")

    _sub(doc, "Scramble suspects - zones with |actual_mean shift| > 40%")
    suspects = patterns["pattern_3_per_zone_divergence"]["scramble_suspects"]
    rows = []
    for r in suspects:
        rows.append([r["zone"], r["borough"], r["actual_mean_baseline"],
                     r["actual_mean_new"], f"{r['actual_mean_drop']*100:+.1f}%"])
    _table(doc,
        ["Zone", "Borough id", "Baseline actual mean", "Drift actual mean", "Change"],
        rows)

    _commentary(doc,
                "Explain the Brooklyn paradox in plain English: why a 'better' "
                "deviance number can hide drift. Note which zones are most "
                "suspect (especially outer-borough ones, Queens and Brooklyn) "
                "and the hypothesized scramble mechanism.")

    # ------------------------------------------------------------------
    # Pattern 4
    # ------------------------------------------------------------------
    _section(doc, "Pattern 4: Manhattan Weekend Concept Drift")
    _table(doc,
        ["Field", "Value"],
        [
            ["What drifted",
             "feature to trip_count mapping for Manhattan x weekend rows"],
            ["Type",
             "Concept drift (input distribution unchanged; target relationship changed)"],
            ["Detection layer",
             "metric_2 per (borough x is_weekend)"],
            ["Why this is hard",
             "Cannot be caught from input-side distributional tests alone. "
             "Needs prediction-vs-actual comparison."],
        ])

    _sub(doc, "Per (borough x is_weekend) - Poisson deviance and MAPE changes")
    pbw = metrics["metrics"]["metric_2_accuracy_by_zone"]["per_borough_weekend"]
    rows = []
    for key in sorted(pbw.keys()):
        s = pbw[key]
        b_part, w_part = key.split("_")
        bid = int(b_part[1:]); w = int(w_part[1:])
        dev_d = s.get("deviance_rel_delta")
        mape_d = s.get("mape_abs_delta")
        label = f"{BOROUGH_NAMES.get(bid, '?')} / {'weekend' if w else 'weekday'}"
        rows.append([
            label, s["n"], s["poisson_deviance"],
            f"{dev_d*100:+.1f}%" if dev_d is not None else "-",
            f"{mape_d*100:+.2f} pp" if mape_d is not None else "-",
        ])
    _table(doc,
        ["Segment", "n rows", "Deviance (drift)", "Change in deviance", "Change in MAPE (pp)"],
        rows)

    _commentary(doc,
                "Highlight the Manhattan x weekend row: it is the hotspot. "
                "Contrast versus Manhattan x weekday (much less drift) and "
                "versus other boroughs. State the real-world hypothesis "
                "(weekend tourism reduction / hybrid-work weekend leisure "
                "shift / etc).")

    # ------------------------------------------------------------------
    # Cross-cutting findings
    # ------------------------------------------------------------------
    _section(doc, "Cross-Cutting Findings")

    _sub(doc, "RMSE was the wrong primary accuracy metric")
    _para(doc,
          "The Week 2 model was trained with objective=poisson. "
          "On the Feb 2 to 28 drift window:")
    _table(doc,
        ["Metric", "Baseline", "Drift", "Relative change"],
        [
            ["Poisson deviance",
             m1.get("baseline_poisson_deviance"),
             m1["poisson_deviance"],
             f"{m1['deviance_rel_delta']*100:+.1f}%"],
            ["MAPE (absolute)", "-", m1["mape"], f"{m1['mape_abs_delta']:+.4f} pp"],
            ["RMSE", "-", m1["rmse"], f"{m1['rmse_rel_delta']*100:+.1f}%"],
        ])
    _commentary(doc,
                "Walk the reader through why RMSE is misleading for Poisson "
                "targets. Recommend the design fix: Poisson deviance is the "
                "model's own training loss and is the right primary signal.")

    _sub(doc, "Per-segment metrics are required, not optional")
    _commentary(doc,
                "Reference the Manhattan weekend example: global metrics "
                "wash out because Manhattan weekday rows dilute Manhattan "
                "weekend rows. Without the (borough x is_weekend) cut, "
                "Pattern 4 is structurally invisible.")

    # ------------------------------------------------------------------
    # Manifest self-check
    # ------------------------------------------------------------------
    _section(doc, "Self-Check Against the Planted Manifest")
    _para(doc,
          "The course manifest (week4_drift_manifest.json) documents the "
          "patterns the TA planted in the Feb 2 to 28 window. It was kept "
          "locally as a verification artifact only; no metric was designed "
          "by reading the manifest first.")
    _table(doc,
        ["Manifest pattern", "Detected by", "Match"],
        [
            ["temporal_peak_shift",
             "metric_4 per-hour KS (hours 9 to 10 stand out)",
             "yes"],
            ["manhattan_lag_deflation",
             "metric_5 per-borough PSI on lag features",
             "yes"],
            ["outer_borough_baseline_scramble",
             "Pattern-3 scramble suspects + Brooklyn paradox",
             "yes"],
            ["manhattan_weekend_concept_drift",
             "metric_2 per (borough x is_weekend)",
             "yes"],
        ])
    _commentary(doc,
                "Optional but valuable: 1 to 2 sentences on what the "
                "self-check validates (the framework caught what was "
                "planted) and what it does NOT validate (real-world drift "
                "may differ from the planted patterns).")

    # ------------------------------------------------------------------
    # Appendices
    # ------------------------------------------------------------------
    _section(doc, "Appendix A: Alert scoreboard from compute_metrics run")
    by_metric = {}
    for a in metrics["alerts"]["critical"]:
        by_metric[a["metric"]] = by_metric.get(a["metric"], 0) + 1
    rows = [[m, n] for m, n in sorted(by_metric.items(), key=lambda kv: -kv[1])]
    _table(doc, ["Metric", "Critical alerts fired"], rows)
    _para(doc,
          f"Total CRITICAL: {metrics['alerts']['critical_count']}. "
          f"Total WARNING: {metrics['alerts']['warning_count']}. "
          f"Decision: {metrics['decision']}.")

    _section(doc, "Appendix B: Headline metric_1 detail")
    _table(doc,
        ["Field", "Value"],
        [[k, m1[k]] for k in (
            "n", "poisson_deviance", "rmse", "mape", "actual_mean", "pred_mean",
            "baseline_poisson_deviance", "deviance_rel_delta", "rmse_rel_delta",
            "mape_abs_delta",
        ) if k in m1])

    doc.save(str(OUT))
    print(f"wrote {OUT.relative_to(REPO)}")


if __name__ == "__main__":
    build()
